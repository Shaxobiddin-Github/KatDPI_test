# Foydalanish qo‘llanmasi sahifasi (controller uchun)

# Importlarni yuqoriga ko'chirish
from main.models import GroupSubject, Semester, Group, Bulim, Kafedra, Subject, University, Faculty
# AJAX orqali guruhga tegishli fanlarni qaytaruvchi endpoint
from django.views.decorators.http import require_GET
@require_GET
def get_subjects_by_group(request):
    group_id = request.GET.get('group_id')
    bulim_id = request.GET.get('bulim_id')
    kafedra_id = request.GET.get('kafedra_id')
    subjects = []
    if group_id:
        # Support comma-separated multiple group ids
        group_ids = [gid for gid in str(group_id).split(',') if gid]
        qs = GroupSubject.objects.filter(group_id__in=group_ids)
        # Dedupe by (subject, semester) across groups
        seen = set()
        for gs in qs.select_related('subject', 'semester'):
            key = (gs.subject_id, gs.semester_id)
            if key in seen:
                continue
            seen.add(key)
            subjects.append({
                'id': gs.subject.id,
                'name': gs.subject.name,
                'semester': gs.semester.number if gs.semester else None,
                'semester_id': gs.semester.id if gs.semester else None,
            })
    elif bulim_id:
        bulim_subjects = GroupSubject.objects.filter(bulim_id=bulim_id)
        for gs in bulim_subjects.select_related('subject'):
            subjects.append({
                'id': gs.subject.id,
                'name': gs.subject.name,
                'semester': None,
                'semester_id': None,
            })
    elif kafedra_id:
        kafedra_subjects = GroupSubject.objects.filter(kafedra_id=kafedra_id)
        for gs in kafedra_subjects.select_related('subject'):
            subjects.append({
                'id': gs.subject.id,
                'name': gs.subject.name,
                'semester': None,
                'semester_id': None,
            })
    return JsonResponse({'subjects': subjects})
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import io
import openpyxl
from django.http import HttpResponse
from main.models import User, Kafedra, Bulim
from django.views.decorators.http import require_POST
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, redirect
from main.models import Test, Subject, Question, Group
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.paginator import Paginator
# Faqat controller uchun dekorator
def controller_required(view_func):
    return user_passes_test(lambda u: u.is_authenticated and u.role == 'controller')(view_func)


@controller_required
def controller_help(request):
    return render(request, 'controller_panel/help.html')
# --- WORD EXPORT ---

# GroupSubject ro'yxati (faqat controller)
@controller_required


@controller_required
def group_subjects_list(request):

    # Universitet o'chirish
    delete_university_id = request.GET.get('delete_university')
    if delete_university_id:
        try:
            University.objects.filter(id=delete_university_id).delete()
            msg = 'Universitet o‘chirildi.'
        except Exception as e:
            msg = f'Universitet o‘chirishda xatolik: {e}'

    # Universitet tahrirlash
    edit_university_id = request.GET.get('edit_university')
    edit_university = None
    if edit_university_id:
        try:
            edit_university = University.objects.get(id=edit_university_id)
        except University.DoesNotExist:
            edit_university = None
            msg = 'Universitet topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_university_id') == edit_university_id:
            edit_university.name = request.POST.get('university_name')
            try:
                edit_university.full_clean()
                edit_university.save()
                msg = 'Universitet tahrirlandi.'
                edit_university = None
            except Exception as e:
                msg = f'Universitet tahrirda xatolik: {e}'

    # Fakultet o'chirish
    delete_faculty_id = request.GET.get('delete_faculty')
    if delete_faculty_id:
        try:
            Faculty.objects.filter(id=delete_faculty_id).delete()
            msg = 'Fakultet o‘chirildi.'
        except Exception as e:
            msg = f'Fakultet o‘chirishda xatolik: {e}'

    # Fakultet tahrirlash
    edit_faculty_id = request.GET.get('edit_faculty')
    edit_faculty = None
    if edit_faculty_id:
        try:
            edit_faculty = Faculty.objects.get(id=edit_faculty_id)
        except Faculty.DoesNotExist:
            edit_faculty = None
            msg = 'Fakultet topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_faculty_id') == edit_faculty_id:
            edit_faculty.name = request.POST.get('faculty_name')
            edit_faculty.university_id = request.POST.get('university')
            try:
                edit_faculty.full_clean()
                edit_faculty.save()
                msg = 'Fakultet tahrirlandi.'
                edit_faculty = None
            except Exception as e:
                msg = f'Fakultet tahrirda xatolik: {e}'

    # Bo'lim o'chirish
    delete_bulim_id = request.GET.get('delete_bulim')
    if delete_bulim_id:
        try:
            Bulim.objects.filter(id=delete_bulim_id).delete()
            msg = 'Bo‘lim o‘chirildi.'
        except Exception as e:
            msg = f'Bo‘lim o‘chirishda xatolik: {e}'

    # Bo'lim tahrirlash
    edit_bulim_id = request.GET.get('edit_bulim')
    edit_bulim = None
    if edit_bulim_id:
        try:
            edit_bulim = Bulim.objects.get(id=edit_bulim_id)
        except Bulim.DoesNotExist:
            edit_bulim = None
            msg = 'Bo‘lim topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_bulim_id') == edit_bulim_id:
            edit_bulim.name = request.POST.get('bulim_name')
            try:
                edit_bulim.full_clean()
                edit_bulim.save()
                msg = 'Bo‘lim tahrirlandi.'
                edit_bulim = None
            except Exception as e:
                msg = f'Bo‘lim tahrirda xatolik: {e}'

    # Guruh o'chirish
    delete_group_id = request.GET.get('delete_group')
    if delete_group_id:
        try:
            Group.objects.filter(id=delete_group_id).delete()
            msg = 'Guruh o‘chirildi.'
        except Exception as e:
            msg = f'Guruh o‘chirishda xatolik: {e}'

    # Guruh tahrirlash
    edit_group_id = request.GET.get('edit_group')
    edit_group = None
    if edit_group_id:
        try:
            edit_group = Group.objects.get(id=edit_group_id)
        except Group.DoesNotExist:
            edit_group = None
            msg = 'Guruh topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_group_id') == edit_group_id:
            edit_group.name = request.POST.get('group_name')
            edit_group.faculty_id = request.POST.get('faculty')
            try:
                edit_group.full_clean()
                edit_group.save()
                msg = 'Guruh tahrirlandi.'
                edit_group = None
            except Exception as e:
                msg = f'Guruh tahrirda xatolik: {e}'

    # Fan o'chirish
    delete_subject_id = request.GET.get('delete_subject')
    if delete_subject_id:
        try:
            Subject.objects.filter(id=delete_subject_id).delete()
            msg = 'Fan o‘chirildi.'
        except Exception as e:
            msg = f'Fan o‘chirishda xatolik: {e}'

    # Fan tahrirlash
    edit_subject_id = request.GET.get('edit_subject')
    edit_subject = None
    if edit_subject_id:
        try:
            edit_subject = Subject.objects.get(id=edit_subject_id)
        except Subject.DoesNotExist:
            edit_subject = None
            msg = 'Fan topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_subject_id') == edit_subject_id:
            edit_subject.name = request.POST.get('subject_name')
            try:
                edit_subject.full_clean()
                edit_subject.save()
                msg = 'Fan tahrirlandi.'
                edit_subject = None
            except Exception as e:
                msg = f'Fan tahrirda xatolik: {e}'

    # Kafedra o'chirish
    delete_kafedra_id = request.GET.get('delete_kafedra')
    if delete_kafedra_id:
        try:
            Kafedra.objects.filter(id=delete_kafedra_id).delete()
            msg = 'Kafedra o‘chirildi.'
        except Exception as e:
            msg = f'Kafedra o‘chirishda xatolik: {e}'

    # Kafedra tahrirlash (GET: formni to'ldirish, POST: saqlash)
    edit_kafedra_id = request.GET.get('edit_kafedra')
    edit_kafedra = None
    if edit_kafedra_id:
        try:
            edit_kafedra = Kafedra.objects.get(id=edit_kafedra_id)
        except Kafedra.DoesNotExist:
            edit_kafedra = None
            msg = 'Kafedra topilmadi.'
        # Tahrirlashni saqlash
        if request.method == 'POST' and request.POST.get('edit_kafedra_id') == edit_kafedra_id:
            edit_kafedra.name = request.POST.get('kafedra_name')
            edit_kafedra.faculty_id = request.POST.get('faculty')
            try:
                edit_kafedra.full_clean()
                edit_kafedra.save()
                msg = 'Kafedra tahrirlandi.'
                edit_kafedra = None
            except Exception as e:
                msg = f'Kafedra tahrirda xatolik: {e}'
    msg = ''
    edit_gs = None

    # O'chirish
    delete_id = request.GET.get('delete')
    if delete_id:
        try:
            GroupSubject.objects.filter(id=delete_id).delete()
            msg = 'Bog‘lash o‘chirildi.'
        except Exception as e:
            msg = f'O‘chirishda xatolik: {e}'



    # Yangi universitet qo'shish
    if request.method == 'POST' and 'add_university' in request.GET:
        name = request.POST.get('university_name')
        if name:
            University.objects.create(name=name)
            msg = 'Universitet qo‘shildi.'

    # Yangi fakultet qo'shish
    if request.method == 'POST' and 'add_faculty' in request.GET:
        name = request.POST.get('faculty_name')
        university_id = request.POST.get('university')
        if name and university_id:
            Faculty.objects.create(name=name, university_id=university_id)
            msg = 'Fakultet qo‘shildi.'

    # Yangi kafedra qo'shish (fakultetdan keyin)
    if request.method == 'POST' and 'add_kafedra' in request.GET:
        name = request.POST.get('kafedra_name')
        faculty_id = request.POST.get('faculty')
        if name and faculty_id:
            Kafedra.objects.create(name=name, faculty_id=faculty_id)
            msg = 'Kafedra qo‘shildi.'

    # Yangi bo'lim qo'shish
    if request.method == 'POST' and 'add_bulim' in request.GET:
        name = request.POST.get('bulim_name')
        if name:
            Bulim.objects.create(name=name)
            msg = 'Bo‘lim qo‘shildi.'

    # Yangi guruh qo'shish
    if request.method == 'POST' and 'add_group' in request.GET:
        name = request.POST.get('group_name')
        faculty_id = request.POST.get('faculty')
        if name and faculty_id:
            Group.objects.create(name=name, faculty_id=faculty_id)
            msg = 'Guruh qo‘shildi.'

    # Yangi fan qo'shish
    if request.method == 'POST' and 'add_subject' in request.GET:
        name = request.POST.get('subject_name')
        if name:
            Subject.objects.create(name=name)
            msg = 'Fan qo‘shildi.'

    # Tahrirlash (GET: formni to'ldirish, POST: saqlash)
    edit_id = request.GET.get('edit')
    if edit_id:
        try:
            edit_gs = GroupSubject.objects.get(id=edit_id)
        except GroupSubject.DoesNotExist:
            edit_gs = None
            msg = 'Bog‘lash topilmadi.'
        # Tahrirlashni saqlash
        if request.method == 'POST' and request.POST.get('edit_id') == edit_id:
            edit_gs.group_id = request.POST.get('group') or None
            edit_gs.bulim_id = request.POST.get('bulim') or None
            edit_gs.kafedra_id = request.POST.get('kafedra') or None
            edit_gs.subject_id = request.POST.get('subject')
            edit_gs.semester_id = request.POST.get('semester') or None
            try:
                edit_gs.full_clean()
                edit_gs.save()
                msg = 'Tahrirlandi.'
                edit_gs = None
            except Exception as e:
                msg = f'Tahrirda xatolik: {e}'
    # GroupSubject qo'shish (faqat shu form yuborilganda)
    elif request.method == 'POST' and request.POST.get('add_group_subject') == '1':
        group_id = request.POST.get('group') or None
        bulim_id = request.POST.get('bulim') or None
        kafedra_id = request.POST.get('kafedra') or None
        # Subject may include semester in value as "<subject_id>::<semester_id>" for student target
        raw_subject = request.POST.get('subject')
        subject_id = None
        semester_id_for_test = None
        if raw_subject and '::' in raw_subject:
            try:
                parts = raw_subject.split('::', 1)
                subject_id = parts[0]
                semester_id_for_test = parts[1] or None
            except Exception:
                subject_id = raw_subject
        else:
            subject_id = raw_subject
        semester_id = request.POST.get('semester') or None
        gs = GroupSubject(
            group_id=group_id if group_id else None,
            bulim_id=bulim_id if bulim_id else None,
            kafedra_id=kafedra_id if kafedra_id else None,
            subject_id=subject_id,
            semester_id=semester_id if semester_id else None
        )
        try:
            gs.full_clean()
            gs.save()
            msg = 'Bog‘lash muvaffaqiyatli qo‘shildi.'
        except Exception as e:
            msg = f'Xatolik: {e}'

    # Pagination for GroupSubjects
    group_subjects_list = GroupSubject.objects.select_related('group', 'bulim', 'kafedra', 'subject', 'semester').all()
    paginator = Paginator(group_subjects_list, 20)  # 20 items per page
    page_number = request.GET.get('page')
    group_subjects = paginator.get_page(page_number)
    
    # Pagination for entities
    universities_list = University.objects.all()
    uni_paginator = Paginator(universities_list, 10)
    uni_page = request.GET.get('uni_page')
    universities = uni_paginator.get_page(uni_page)
    
    faculties_list = Faculty.objects.select_related('university').all()
    fac_paginator = Paginator(faculties_list, 10)
    fac_page = request.GET.get('fac_page')
    faculties = fac_paginator.get_page(fac_page)
    
    kafedra_list = Kafedra.objects.select_related('faculty').all()
    kaf_paginator = Paginator(kafedra_list, 10)
    kaf_page = request.GET.get('kaf_page')
    kafedralar = kaf_paginator.get_page(kaf_page)
    
    bulim_list = Bulim.objects.all()
    bul_paginator = Paginator(bulim_list, 10)
    bul_page = request.GET.get('bul_page')
    bulimlar = bul_paginator.get_page(bul_page)
    
    group_list = Group.objects.select_related('faculty').all()
    grp_paginator = Paginator(group_list, 10)
    grp_page = request.GET.get('grp_page')
    groups = grp_paginator.get_page(grp_page)
    
    subject_list = Subject.objects.all()
    sub_paginator = Paginator(subject_list, 10)
    sub_page = request.GET.get('sub_page')
    subjects = sub_paginator.get_page(sub_page)

    context = {
        'group_subjects': group_subjects,
        'groups': groups,
        'bulimlar': bulimlar,
        'kafedralar': kafedralar,
        'subjects': subjects,
        'semesters': Semester.objects.all(),
        'universities': universities,
        'faculties': faculties,
        'all_groups': Group.objects.all(),  # For select dropdowns
        'all_bulimlar': Bulim.objects.all(),
        'all_kafedralar': Kafedra.objects.all(),
        'all_subjects': Subject.objects.all(),
        'all_universities': University.objects.all(),
        'all_faculties': Faculty.objects.all(),
        'msg': msg,
        'edit_gs': edit_gs,
        'edit_kafedra': edit_kafedra,
        'edit_university': edit_university,
        'edit_faculty': edit_faculty,
        'edit_bulim': edit_bulim,
        'edit_group': edit_group,
        'edit_subject': edit_subject,
    }
    return render(request, 'controller_panel/group_subjects.html', context)
from django.utils import timezone

# --- WORD EXPORT ---
@login_required
def export_users_word(request):
    users = User.objects.exclude(is_superuser=True)
    # Filterlar
    filter_role = request.GET.get('filter_role')
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_role:
        users = users.filter(role=filter_role)
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)
    # Role bo‘yicha ajratish
    role_map = defaultdict(list)
    for user in users:
        role_map[user.get_role_display()].append(user)
    doc = Document()
    doc.add_heading('Foydalanuvchilar ro‘yxati', 0)
    for role, userlist in role_map.items():
        doc.add_heading(role, level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ism'
        hdr_cells[1].text = 'Familiya'
        hdr_cells[2].text = 'Username'
        hdr_cells[3].text = 'Guruh/Kafedra/Bo‘lim'
        hdr_cells[4].text = 'Access code'
        hdr_cells[5].text = 'Role'
        seen = set()
        for u in userlist:
            key = (u.first_name, u.last_name, u.username, u.role, u.group_id, u.kafedra_id, u.bulim_id)
            if key in seen:
                continue
            seen.add(key)
            row = table.add_row().cells
            row[0].text = u.first_name or ''
            row[1].text = u.last_name or ''
            row[2].text = u.username or ''
            # Guruh/kafedra/bulim bir ustunda
            if u.role == 'student' and u.group:
                row[3].text = u.group.name
            elif u.role == 'tutor' and u.kafedra:
                row[3].text = u.kafedra.name
            elif u.role == 'employee' and u.bulim:
                row[3].text = u.bulim.name
            else:
                row[3].text = ''
            row[4].text = u.access_code or ''
            row[5].text = u.get_role_display()
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=foydalanuvchilar.docx'
    return response

# --- PDF EXPORT ---
@login_required
def export_users_pdf(request):
    users = User.objects.exclude(is_superuser=True)
    filter_role = request.GET.get('filter_role')
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_role:
        users = users.filter(role=filter_role)
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)
    role_map = defaultdict(list)
    for user in users:
        role_map[user.get_role_display()].append(user)
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    y = height - 40
    c.setFont('Helvetica-Bold', 16)
    c.drawString(40, y, 'Foydalanuvchilar ro‘yxati')
    y -= 30
    c.setFont('Helvetica', 12)
    for role, userlist in role_map.items():
        c.setFont('Helvetica-Bold', 14)
        c.drawString(40, y, role)
        y -= 22
        c.setFont('Helvetica', 10)
        c.drawString(40, y, 'Ism')
        c.drawString(120, y, 'Familiya')
        c.drawString(220, y, 'Username')
        c.drawString(320, y, 'Guruh/Kafedra/Bo‘lim')
        c.drawString(470, y, 'Access code')
        y -= 16
        seen = set()
        for u in userlist:
            key = (u.first_name, u.last_name, u.username, u.role, u.group_id, u.kafedra_id, u.bulim_id)
            if key in seen:
                continue
            seen.add(key)
            c.drawString(40, y, u.first_name or '')
            c.drawString(120, y, u.last_name or '')
            c.drawString(220, y, u.username or '')
            if u.role == 'student' and u.group:
                c.drawString(320, y, u.group.name)
            elif u.role == 'tutor' and u.kafedra:
                c.drawString(320, y, u.kafedra.name)
            elif u.role == 'employee' and u.bulim:
                c.drawString(320, y, u.bulim.name)
            else:
                c.drawString(320, y, '')
            c.drawString(470, y, u.access_code or '')
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40
                c.setFont('Helvetica', 10)
        y -= 18
    c.save()
    output.seek(0)
    response = HttpResponse(output, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=foydalanuvchilar.pdf'
    return response
@login_required
def export_users_excel(request):
    filter_role = request.GET.get('filter_role')
    users = User.objects.exclude(is_superuser=True)
    if filter_role:
        users = users.filter(role=filter_role)
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Foydalanuvchilar'
    ws.append(['Ism', 'Familiya', 'Access code'])
    for user in users:
        ws.append([
            getattr(user, 'first_name', ''),
            getattr(user, 'last_name', ''),
            getattr(user, 'access_code', '')
        ])
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=foydalanuvchilar.xlsx'
    return response


# Talabalar importi uchun Excel shablonini yuklab berish
@login_required
def download_student_import_template(request):
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Talaba Import'
    # Moslashuvchanlik uchun asosiy sarlavhalarni beramiz (role kerak emas, hammasi student)
    ws.append(['firstname', 'lastname', 'group'])
    # Namuna qatorlar
    ws.append(['Ali', 'Valiyev', 'CS-101'])
    ws.append(['Laylo', 'Karimova', 'CS-102'])
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=talaba_import_shablon.xlsx'
    return response


# Foydalanuvchilarni ko‘rish va qo‘shish (superuserlarsiz)
@login_required
def add_user(request):
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
    users = User.objects.exclude(is_superuser=True)
    # Filter GET params
    filter_role = request.GET.get('filter_role')
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_role:
        users = users.filter(role=filter_role)
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)

    groups = Group.objects.all()
    kafedralar = Kafedra.objects.all()
    bulimlar = Bulim.objects.all()
    role_choices = User.ROLE_CHOICES
    if request.method == 'POST':
        # Branch 1: Excel import for students
        if request.POST.get('import_type') == 'student' and request.FILES.get('excel_file'):
            from django.db import transaction
            import random
            import string
            file = request.FILES['excel_file']
            wb = openpyxl.load_workbook(file)
            ws = wb.active
            headers = {}
            results = []  # for display: list of dicts
            errors = []

            # Normalize header row
            header_row = None
            for row in ws.iter_rows(min_row=1, max_row=1, values_only=True):
                header_row = [str(c).strip().lower() if c is not None else '' for c in row]
                break
            if not header_row:
                request.session['import_result'] = {
                    'errors': ['Excel fayl bo\'sh yoki sarlavha (Header) yo\'q'],
                }
                return redirect('add_user')

            # Map possible header names
            def find_idx(names):
                for name in names:
                    if name in header_row:
                        return header_row.index(name)
                return None

            idx_first = find_idx(['firstname', 'first_name', 'ism', 'first name'])
            idx_last = find_idx(['lastname', 'last_name', 'familiya', 'last name'])
            idx_group = find_idx(['group', 'guruh', 'group_name'])

            if idx_first is None or idx_last is None:
                request.session['import_result'] = {
                    'errors': ["Sarlavha topilmadi. Kerakli ustunlar: firstname, lastname (ixtiyoriy: group)"],
                }
                return redirect('add_user')

            # Helper: generate password
            def gen_password(length=8):
                chars = string.ascii_letters + string.digits
                return ''.join(random.choice(chars) for _ in range(length))

            # Helper: ensure unique username
            def make_unique_username(base):
                base_clean = (base or 'user').strip().lower().replace(' ', '')
                if not base_clean:
                    base_clean = 'user'
                candidate = base_clean
                i = 1
                while User.objects.filter(username=candidate).exists():
                    i += 1
                    candidate = f"{base_clean}{i}"
                return candidate

            # Build group name -> id map (case-insensitive)
            group_by_name = {g.name.strip().lower(): g.id for g in groups}

            created_count = 0
            with transaction.atomic():
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if not row:
                        continue
                    first = str(row[idx_first]).strip() if idx_first is not None and row[idx_first] is not None else ''
                    last = str(row[idx_last]).strip() if idx_last is not None and row[idx_last] is not None else ''
                    if not first and not last:
                        # skip empty
                        continue
                    group_name = str(row[idx_group]).strip() if idx_group is not None and row[idx_group] is not None else ''
                    # Roli qatordan qat'i nazar, barcha yaratiladigan foydalanuvchilar student bo'ladi

                    username = make_unique_username(first)
                    password = gen_password()
                    user = User(username=username, role='student', first_name=first, last_name=last)
                    # Assign group by name
                    if group_name:
                        gid = group_by_name.get(group_name.strip().lower())
                        if gid:
                            user.group_id = gid
                        else:
                            errors.append(f"{first} {last}: Guruh topilmadi: {group_name}")
                    user.set_password(password)
                    user.save()  # access_code auto-generates in model save

                    created_count += 1
                    results.append({
                        'full_name': f"{first} {last}".strip(),
                        'username': username,
                        'password': password,
                        'access_code': user.access_code or '',
                        'group': group_name or '-',
                    })

            request.session['import_result'] = {
                'created': created_count,
                'errors': errors,
                'rows': results,
            }
            return redirect('add_user')

        # Branch 2: Single user create (existing behavior)
        username = request.POST.get('username')
        password = request.POST.get('password')
        role = request.POST.get('role')
        group_id = request.POST.get('group')
        kafedra_id = request.POST.get('kafedra')
        bulim_id = request.POST.get('bulim')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        user = User(username=username, role=role, first_name=first_name, last_name=last_name)
        if group_id:
            user.group_id = group_id
        if kafedra_id:
            user.kafedra_id = kafedra_id
        if bulim_id:
            user.bulim_id = bulim_id
        user.set_password(password)
        user.save()
        return redirect('add_user')
    return render(request, 'controller_panel/add_user.html', {
        'users': users,
        'groups': groups,
        'kafedralar': kafedralar,
        'bulimlar': bulimlar,
        'role_choices': role_choices,
    'import_result': request.session.pop('import_result', None),
    })
@login_required
def controller_logout(request):
    logout(request)
    return redirect('/api/login/')
# AJAX: Controller o‘zining savolini o‘chira oladi (lekin tahrirlay olmaydi)

def login_check(request):
    if not request.user.is_authenticated:
        return redirect('/api/login/')
    return None

@login_required
@require_POST
def delete_question(request, question_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return JsonResponse({'success': False, 'error': 'Ruxsat yo‘q'}, status=403)
    deleted, _ = Question.objects.filter(id=question_id).delete()
    return JsonResponse({'success': bool(deleted)})


# AJAX: Fanga tegishli savollar ro‘yxati (faqat controller o‘zining savollarini ko‘radi)
from django.views.decorators.http import require_GET
@login_required
@require_GET
def subject_questions(request, subject_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return render(request, 'controller_panel/questions_by_subject.html', {'questions': [], 'subject': {'name': 'Noma’lum'}})
    from main.models import Subject
    subject = Subject.objects.filter(id=subject_id).first()
    questions = Question.objects.filter(subject_id=subject_id)
    return render(request, 'controller_panel/questions_by_subject.html', {
        'questions': questions,
        'subject': subject
    })

# Testni tahrirlash (edit)
@login_required
def edit_test(request, test_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
    from main.models import TestQuestion
    test = Test.objects.get(id=test_id, created_by=request.user)
    groups = Group.objects.all()
    kafedralar = Kafedra.objects.all()
    bulimlar = Bulim.objects.all()
    subjects = Subject.objects.all()
    if request.method == 'POST':
        raw_subject = request.POST.get('subject')
        subject_id = None
        semester_id_for_test = None
        # POST'd subject may arrive as "<subject_id>::<semester_id>" (student/group-target flow).
        # Parse robustly without relying on a separate 'target' flag.
        if raw_subject:
            if '::' in raw_subject:
                subject_id, semester_id_for_test = raw_subject.split('::', 1)
            else:
                subject_id = raw_subject
        question_count = int(request.POST.get('question_count'))
        total_score = int(request.POST.get('total_score'))
        duration_str = request.POST.get('duration')
        video_url = (request.POST.get('video_url') or '').strip()
        video_file = request.FILES.get('video_file')
        # YouTube URL bo'lsa, embed formatga aylantiramiz
        def to_youtube_embed(url: str) -> str:
            if not url:
                return url
            u = url.strip()
            # short youtu.be/VIDEOID
            if 'youtu.be/' in u:
                vid = u.split('youtu.be/', 1)[1].split('?')[0].strip('/')
                return f"https://www.youtube.com/embed/{vid}"
            # watch?v=VIDEOID
            if 'watch?v=' in u:
                vid = u.split('watch?v=', 1)[1].split('&')[0]
                return f"https://www.youtube.com/embed/{vid}"
            # already embed
            if '/embed/' in u:
                return u
            return u
        video_url = to_youtube_embed(video_url)
        duration = parse_duration_string(duration_str)

        if duration is None:
            return render(request, 'controller_panel/edit_test.html', {
                'groups': groups,
                'kafedralar': kafedralar,
                'bulimlar': bulimlar,
                'subjects': subjects,
                'test': test,
                'error': 'Test muddati noto‘g‘ri formatda! To‘g‘ri format: soat:daqiq:soniya (masalan: 00:30:00)'
            })
        subject = Subject.objects.get(id=subject_id)
        # Test turi bo‘yicha mos maydonlarni yangilash
        if test.group:
            group_id = request.POST.get('group')
            group = Group.objects.get(id=group_id)
            test.group = group
            test.kafedra = None
            test.bulim = None
        elif test.kafedra:
            kafedra_id = request.POST.get('kafedra')
            kafedra = Kafedra.objects.get(id=kafedra_id)
            test.kafedra = kafedra
            test.group = None
            test.bulim = None
        elif test.bulim:
            bulim_id = request.POST.get('bulim')
            bulim = Bulim.objects.get(id=bulim_id)
            test.bulim = bulim
            test.group = None
            test.kafedra = None
        test.subject = subject
        # If semester was provided in the subject selector, update it; otherwise preserve existing value.
        if semester_id_for_test:
            try:
                test.semester_id = int(semester_id_for_test)
            except (TypeError, ValueError):
                test.semester_id = semester_id_for_test  # fallback: let ORM coerce if it's a numeric string
        test.question_count = question_count
        test.total_score = total_score
        test.duration = duration
        # Update optional video (remove if requested)
        if request.POST.get('remove_video') == '1':
            test.video_url = None
            if getattr(test, 'video_file', None):
                test.video_file.delete(save=False)
            test.video_file = None
        else:
            test.video_url = video_url or None
            if video_file:
                test.video_file = video_file
        test.start_time = timezone.now()
        test.save()
        # Savollarni yangilash (oddiy variant: eski TestQuestionlarni o‘chirib, yangidan yaratamiz)
        TestQuestion.objects.filter(test=test).delete()
        # Filter questions by subject and semester if test has semester
        if getattr(test, 'semester_id', None):
            questions = Question.objects.filter(subject=subject, semester_id=test.semester_id)
        else:
            questions = Question.objects.filter(subject=subject)
        if questions.count() < question_count:
            return render(request, 'controller_panel/edit_test.html', {
                'groups': groups,
                'kafedralar': kafedralar,
                'bulimlar': bulimlar,
                'subjects': subjects,
                'test': test,
                'error': 'Ushbu fanga yetarli savol mavjud emas!'
            })
        selected_questions = questions.order_by('?')[:question_count]
        score_per_question = total_score / question_count
        for question in selected_questions:
            TestQuestion.objects.create(test=test, question=question, score=score_per_question)
        return redirect('controller_dashboard')
    # GET
    return render(request, 'controller_panel/edit_test.html', {
        'groups': groups,
        'kafedralar': kafedralar,
        'bulimlar': bulimlar,
        'subjects': subjects,
        'test': test,
        'preset_list': ['00:10:00','00:15:00','00:30:00','00:45:00','01:00:00','01:30:00']
    })

# Testni o‘chirish (delete)
@login_required
@csrf_exempt
def delete_test(request, test_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
    if request.method == 'POST':
        Test.objects.filter(id=test_id, created_by=request.user).delete()
    return redirect('controller_dashboard')


@login_required
def controller_dashboard(request):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
 
    tests = Test.objects.filter(created_by=request.user)
    now = timezone.now()
    # Har bir test uchun status va fanning nomini aniqlash
    test_list = []
    for test in tests:
        status = "Faol"
        if test.end_time and now > test.end_time:
            status = "Test muddati tugagan"
        subject_name = test.subject.name if test.subject else ""
        semester_label = None
        if getattr(test, 'semester_id', None):
            try:
                semester_label = test.semester.number
            except Exception:
                semester_label = None
        test_list.append({
            'test': test,
            'status': status,
            'subject_name': subject_name,
            'semester_label': semester_label,
        })
    # frontendda tests deb ishlatiladi
    return render(request, 'controller_panel/dashboard.html', {'tests': test_list})








from datetime import timedelta

def parse_duration_string(duration_str):
    """Masalan: '00:30:00' ni timedelta ga aylantiradi"""
    try:
        h, m, s = map(int, duration_str.split(":"))
        return timedelta(hours=h, minutes=m, seconds=s)
    except:
        return None


from main.models import Kafedra, Bulim

@login_required
def add_test(request):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')

    target = request.GET.get('target', 'student')
    context = {}
    if target == 'student':
        context['groups'] = Group.objects.all()
        group_id = request.GET.get('group') or request.POST.get('group')
        if group_id:
            from main.models import GroupSubject
            group_subjects = GroupSubject.objects.filter(group_id=group_id)
            context['subjects'] = Subject.objects.filter(id__in=group_subjects.values_list('subject_id', flat=True))
        else:
            context['subjects'] = Subject.objects.none()
        context['target'] = 'student'
    elif target == 'tutor':
        context['kafedralar'] = Kafedra.objects.all()
        context['subjects'] = Subject.objects.all()
        context['target'] = 'tutor'
    elif target == 'employee':
        context['bulimlar'] = Bulim.objects.all()
        context['subjects'] = Subject.objects.all()
        context['target'] = 'employee'
    else:
        context['groups'] = Group.objects.all()
        context['subjects'] = Subject.objects.all()
        context['target'] = 'student'

    if request.method == 'POST':
        raw_subject = request.POST.get('subject')
        subject_id = None
        semester_id_for_test = None
        if raw_subject and target == 'student':
            if '::' in raw_subject:
                subject_id, semester_id_for_test = raw_subject.split('::', 1)
            else:
                subject_id = raw_subject
        else:
            subject_id = raw_subject
        question_count = request.POST.get('question_count')
        total_score = request.POST.get('total_score')
        duration_str = request.POST.get('duration')
        minutes = request.POST.get('minutes', 30)
        # Video maydonlari
        video_url = (request.POST.get('video_url') or '').strip()
        # Fayl yuklash uchun add_test formi enctype multipart bo'lishi kerak; bu yerda POST'dan alohida FILES'dan olamiz
        video_file = request.FILES.get('video_file')
        # Majburiy maydonlar to'ldirilganini tekshirish
        if not (subject_id and question_count and total_score and duration_str):
            context['error'] = "Iltimos, barcha maydonlarni to'ldiring!"
            return render(request, 'controller_panel/add_test.html', context)
        try:
            question_count = int(question_count)
            total_score = int(total_score)
            minutes = int(minutes)
        except ValueError:
            context['error'] = "Sonli maydonlarga faqat raqam kiriting!"
            return render(request, 'controller_panel/add_test.html', context)
        duration = parse_duration_string(duration_str)
        context['error'] = None
        group = None
        kafedra = None
        bulim = None
        if target == 'student':
            # support multiple groups
            group_ids = request.POST.getlist('groups') or []
            if not group_ids:
                context['error'] = "Iltimos, kamida bitta guruhni tanlang!"
                return render(request, 'controller_panel/add_test.html', context)
        elif target == 'tutor':
            kafedra_id = request.POST.get('kafedra')
            kafedra = Kafedra.objects.get(id=kafedra_id)
        elif target == 'employee':
            bulim_id = request.POST.get('bulim')
            bulim = Bulim.objects.get(id=bulim_id)

        if duration is None:
            return render(request, 'controller_panel/add_test.html', context | {'error': 'Test muddati noto‘g‘ri formatda! To‘g‘ri format: soat:daqiq:soniya (masalan: 00:30:00)'})

        subject = Subject.objects.get(id=subject_id)
        from main.models import TestQuestion, Test
        # Test yaratish: modelga mos ravishda group, kafedra yoki bulimni saqlash kerak
        tests = []
        if target == 'student':
            # Create a separate Test for each selected group
            for gid in group_ids:
                try:
                    grp = Group.objects.get(id=gid)
                except Group.DoesNotExist:
                    continue
                tests.append(Test.objects.create(
                    group=grp,
                    subject=subject,
                    semester_id=semester_id_for_test,
                    question_count=question_count,
                    total_score=total_score,
                    duration=duration,
                    minutes=minutes,
                    created_by=request.user,
                    video_url=video_url or None,
                    video_file=video_file if video_file else None
                ))
        elif target == 'tutor':
            # group o‘rniga kafedra saqlash uchun modelda mos o‘zgartirish kerak bo‘ladi
            tests.append(Test.objects.create(
                group=None,
                subject=subject,
                question_count=question_count,
                total_score=total_score,
                duration=duration,
                minutes=minutes,
                created_by=request.user,
                kafedra=kafedra
            ))
        elif target == 'employee':
            tests.append(Test.objects.create(
                group=None,
                subject=subject,
                question_count=question_count,
                total_score=total_score,
                duration=duration,
                minutes=minutes,
                created_by=request.user,
                bulim=bulim
            ))
        else:
            tests.append(Test.objects.create(
                group=group,
                subject=subject,
                question_count=question_count,
                total_score=total_score,
                duration=duration,
                minutes=minutes,
                created_by=request.user,
                video_url=video_url or None,
                video_file=video_file if video_file else None
            ))

        # Savollarni tanlash: semestr bo'yicha filtrlang (agar tanlangan bo'lsa)
        if semester_id_for_test:
            questions = Question.objects.filter(subject=subject, semester_id=semester_id_for_test)
        else:
            questions = Question.objects.filter(subject=subject)
        if questions.count() < question_count:
            context['error'] = 'Ushbu fanga yetarli savol mavjud emas!'
            return render(request, 'controller_panel/add_test.html', context)
        selected_questions = questions.order_by('?')[:question_count]
        score_per_question = total_score / question_count
        for test in tests:
            for question in selected_questions:
                TestQuestion.objects.create(test=test, question=question, score=score_per_question)
        return redirect('controller_dashboard')

    return render(request, 'controller_panel/add_test.html', context)
